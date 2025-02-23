from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Title
doc.add_heading('Software Requirements Specification (SRS)', 0)

# Subtitle
doc.add_heading('iLearnSpace - Attendance Management System using Local Area Network (LAN) using Java GUI', level=1)

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc = [
    "1. Introduction",
    "    1.1 Purpose",
    "    1.2 Scope",
    "    1.3 Definitions, Acronyms, and Abbreviations",
    "    1.4 References",
    "2. Overall Description",
    "    2.1 Product Perspective",
    "    2.2 Product Functions",
    "    2.3 User Classes and Characteristics",
    "    2.4 Operating Environment",
    "    2.5 Design and Implementation Constraints",
    "    2.6 Assumptions and Dependencies",
    "3. Specific Requirements",
    "    3.1 Functional Requirements",
    "    3.2 Non-Functional Requirements",
    "    3.3 Interface Requirements",
    "4. System Features",
    "    4.1 Login Page",
    "    4.2 Admin Page",
    "    4.3 Lecturer Page",
    "    4.4 Student Page",
    "5. Database Design",
    "    5.1 Database Schema",
    "    5.2 Entity-Relationship Diagram",
    "6. Appendices"
]
for item in toc:
    doc.add_paragraph(item)

# Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph('The purpose of this document is to provide a detailed description of the iLearnSpace Attendance Management System. It will explain the system\'s functionality, constraints, and requirements.')

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph('iLearnSpace is a university attendance management system that operates over a Local Area Network (LAN) to ensure secure and real-time attendance tracking. The system includes three user roles: Admin, Lecturer, and Student, each with specific functionalities.')

doc.add_heading('1.3 Definitions, Acronyms, and Abbreviations', level=2)
doc.add_paragraph('- LAN: Local Area Network\n- GUI: Graphical User Interface\n- OTP: One-Time Password')

doc.add_heading('1.4 References', level=2)
doc.add_paragraph('- Project Proposal Document\n- Java Swing/AWT Documentation')

# Overall Description
doc.add_heading('2. Overall Description', level=1)
doc.add_heading('2.1 Product Perspective', level=2)
doc.add_paragraph('iLearnSpace is a standalone application that will be deployed on a university\'s local network. It will use Java Swing/AWT for the GUI and a relational database for data storage.')

doc.add_heading('2.2 Product Functions', level=2)
doc.add_paragraph('- User authentication and role-based access control\n- Attendance tracking with OTP verification\n- User and course management\n- Attendance statistics visualization\n- Notifications and communication')

doc.add_heading('2.3 User Classes and Characteristics', level=2)
doc.add_paragraph('- Admin: Manages users, courses, and oversees the attendance process.\n- Lecturer: Schedules classes, takes attendance, and communicates with students.\n- Student: Views scheduled classes, marks attendance, and checks attendance records.')

doc.add_heading('2.4 Operating Environment', level=2)
doc.add_paragraph('- Operating System: Linux\n- Development Environment: Visual Studio Code\n- Database: MySQL or SQLite')

doc.add_heading('2.5 Design and Implementation Constraints', level=2)
doc.add_paragraph('- The system must operate within a LAN.\n- Attendance marking is restricted to the LAN.')

doc.add_heading('2.6 Assumptions and Dependencies', level=2)
doc.add_paragraph('- Users have access to the LAN.\n- The database server is accessible within the LAN.')

# Specific Requirements
doc.add_heading('3. Specific Requirements', level=1)
doc.add_heading('3.1 Functional Requirements', level=2)
doc.add_paragraph('- Login Page: Authenticate users based on their role.\n- Admin Dashboard: Manage users, courses, and attendance.\n- Lecturer Dashboard: Schedule classes and take attendance.\n- Student Dashboard: View classes and mark attendance.')

doc.add_heading('3.2 Non-Functional Requirements', level=2)
doc.add_paragraph('- Performance: The system should handle up to 1000 concurrent users.\n- Security: User data and attendance records must be securely stored.\n- Usability: The GUI should be intuitive and easy to use.')

doc.add_heading('3.3 Interface Requirements', level=2)
doc.add_paragraph('- User Interface: Java Swing/AWT-based GUI.\n- Database Interface: JDBC for database connectivity.')

# System Features
doc.add_heading('4. System Features', level=1)
doc.add_heading('4.1 Login Page', level=2)
doc.add_paragraph('- Admin Login: Requires Username, Password, and User Type (Admin). OTP verification via email.\n- Lecturer & Student Login: Requires Email and Password.')

doc.add_heading('4.2 Admin Page', level=2)
doc.add_paragraph('- User & Course Management: Add, modify, and delete users and courses.\n- Attendance System: Oversee attendance process with OTP verification.\n- Official Notices: Send notices to users.')

doc.add_heading('4.3 Lecturer Page', level=2)
doc.add_paragraph('- Profile Information: View personal and course information.\n- Class Management: Schedule and manage classes.\n- Attendance Visualization: Display attendance statistics.')

doc.add_heading('4.4 Student Page', level=2)
doc.add_paragraph('- Class Information: View scheduled and ongoing classes.\n- Attendance Marking: Mark attendance using OTP.\n- Attendance Records: View past attendance logs and statistics.')

# Database Design
doc.add_heading('5. Database Design', level=1)
doc.add_heading('5.1 Database Schema', level=2)
doc.add_paragraph('```sql\n-- Drop tables if they already exist (to avoid duplication errors during re-runs)\nDROP TABLE IF EXISTS attendance;\nDROP TABLE IF EXISTS notifications;\nDROP TABLE IF EXISTS courses;\nDROP TABLE IF EXISTS student_course;\nDROP TABLE IF EXISTS lecturer_course;\nDROP TABLE IF EXISTS students;\nDROP TABLE IF EXISTS lecturers;\nDROP TABLE IF EXISTS admins;\nDROP TABLE IF EXISTS users;\n\n-- Create a common users table to store login credentials for Admins, Lecturers, and Students\nCREATE TABLE users (\n    user_id       INT PRIMARY KEY AUTO_INCREMENT,\n    email         VARCHAR(100) UNIQUE NOT NULL,\n    password_hash VARCHAR(255) NOT NULL,\n    user_type     ENUM(\'Admin\', \'Lecturer\', \'Student\') NOT NULL,\n    created_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n);\n\n-- Admin Table (Linked to users table)\nCREATE TABLE admins (\n    admin_id INT PRIMARY KEY AUTO_INCREMENT,\n    user_id  INT UNIQUE NOT NULL,\n    full_name VARCHAR(100) NOT NULL,\n    phone VARCHAR(15),\n    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE\n);\n\n-- Lecturer Table (Linked to users table)\nCREATE TABLE lecturers (\n    lecturer_id INT PRIMARY KEY AUTO_INCREMENT,\n    user_id  INT UNIQUE NOT NULL,\n    full_name VARCHAR(100) NOT NULL,\n    department VARCHAR(100) NOT NULL,\n    phone VARCHAR(15),\n    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE\n);\n\n-- Student Table (Linked to users table)\nCREATE TABLE students (\n    student_id INT PRIMARY KEY AUTO_INCREMENT,\n    user_id  INT UNIQUE NOT NULL,\n    full_name VARCHAR(100) NOT NULL,\n    enrollment_number VARCHAR(50) UNIQUE NOT NULL,\n    department VARCHAR(100) NOT NULL,\n    phone VARCHAR(15),\n    FOREIGN KEY (user_id) REFERENCES users(user_id) ON DELETE CASCADE\n);\n\n-- Course Table (Stores all university courses)\nCREATE TABLE courses (\n    course_id   INT PRIMARY KEY AUTO_INCREMENT,\n    course_name VARCHAR(100) NOT NULL,\n    course_code VARCHAR(20) UNIQUE NOT NULL,\n    created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n);\n\n-- Mapping Table: Many-to-Many Relationship Between Students & Courses\nCREATE TABLE student_course (\n    student_id INT NOT NULL,\n    course_id  INT NOT NULL,\n    PRIMARY KEY (student_id, course_id),\n    FOREIGN KEY (student_id) REFERENCES students(student_id) ON DELETE CASCADE,\n    FOREIGN KEY (course_id) REFERENCES courses(course_id) ON DELETE CASCADE\n);\n\n-- Mapping Table: Many-to-Many Relationship Between Lecturers & Courses\nCREATE TABLE lecturer_course (\n    lecturer_id INT NOT NULL,\n    course_id   INT NOT NULL,\n    PRIMARY KEY (lecturer_id, course_id),\n    FOREIGN KEY (lecturer_id) REFERENCES lecturers(lecturer_id) ON DELETE CASCADE,\n    FOREIGN KEY (course_id) REFERENCES courses(course_id) ON DELETE CASCADE\n);\n\n-- Attendance Table (Stores attendance records)\nCREATE TABLE attendance (\n    attendance_id INT PRIMARY KEY AUTO_INCREMENT,\n    student_id    INT NOT NULL,\n    lecturer_id   INT NOT NULL,\n    course_id     INT NOT NULL,\n    attendance_date DATE NOT NULL,\n    status        ENUM(\'Present\', \'Absent\') NOT NULL,\n    created_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n    FOREIGN KEY (student_id) REFERENCES students(student_id) ON DELETE CASCADE,\n    FOREIGN KEY (lecturer_id) REFERENCES lecturers(lecturer_id) ON DELETE CASCADE,\n    FOREIGN KEY (course_id) REFERENCES courses(course_id) ON DELETE CASCADE\n);\n\n-- Notifications Table (Stores messages sent by Admins or Lecturers)\nCREATE TABLE notifications (\n    notification_id INT PRIMARY KEY AUTO_INCREMENT,\n    sender_id       INT NOT NULL,\n    receiver_id     INT,\n    receiver_type   ENUM(\'Student\', \'Lecturer\', \'All\') NOT NULL,\n    message        TEXT NOT NULL,\n    sent_at        TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n    FOREIGN KEY (sender_id) REFERENCES users(user_id) ON DELETE CASCADE\n);\n```')

doc.add_heading('5.2 Entity-Relationship Diagram', level=2)
doc.add_paragraph('![ER Diagram](path/to/er-diagram.png)')

# Appendices
doc.add_heading('6. Appendices', level=1)
doc.add_paragraph('- Appendix A: Glossary\n- Appendix B: Acronyms\n- Appendix C: References')

# Save the document
doc.save('SRS.docx')